import re
import os
import traceback
import logging
import keyring
from io import BytesIO
from datetime import datetime
from pathlib import Path
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse
from core.models import TimbradoReportRequest, DETNotificationRequest, FisconformeRequest
from core.utils import validar_cnpj, ler_sql, _write_excel_with_format

logger = logging.getLogger("sefin_audit_python")
router = APIRouter(prefix="/api/python", tags=["reports"])

# Get project root from environment or handle it
_PROJETO_DIR = Path(__file__).resolve().parent.parent.parent.parent

@router.post("/reports/timbrado")
async def generate_timbrado_report(request: TimbradoReportRequest):
    """Gera relatório em formato Word com Papel Timbrado SEFIN."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        headers = ["Governo do Estado de Rondônia", "Secretaria de Estado de Finanças - SEFIN", "Coordenadoria da Receita Estadual – CRE", request.orgao]
        for text in headers:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(11)
            if text == request.orgao:
                run.bold = True
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("RELATÓRIO")
        run.bold, run.font.size = True, Pt(14)
        doc.add_paragraph()
        
        fields = [
            ("Razão Social", request.razao_social), ("CNPJ", request.cnpj), ("Inscrição Estadual", request.ie),
            ("Regime de Pagamento", request.regime_pagamento), ("Regime Especial", request.regime_especial),
            ("Atividade Principal", request.atividade_principal), ("Endereço", request.endereco), ("DSF", request.num_dsf),
        ]
        for label, value in fields:
            if value:
                p = doc.add_paragraph()
                run_label = p.add_run(f"{label}: ")
                run_label.bold, run_label.font.size = True, Pt(11)
                p.add_run(value).font.size = Pt(11)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("1. OBJETO")
        run.bold, run.font.size = True, Pt(11)
        doc.add_paragraph(request.objeto).style.font.size = Pt(11)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("2. RELATO")
        run.bold, run.font.size = True, Pt(11)
        if request.relato:
            doc.add_paragraph(request.relato).style.font.size = Pt(11)
        
        if request.itens:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("3. NOTIFICAÇÃO DE INCONSISTÊNCIAS")
            run.bold, run.font.size = True, Pt(11)
            for i, item in enumerate(request.itens, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"Item {i} – {item.get('tipo', '')}")
                run.bold, run.font.size = True, Pt(11)
                if item.get('descricao'):
                    doc.add_paragraph(item['descricao'])
        
        doc.add_paragraph()
        section_num = 4 if request.itens else 3
        p = doc.add_paragraph()
        run = p.add_run(f"{section_num}. CONCLUSÃO")
        run.bold, run.font.size = True, Pt(11)
        if request.conclusao:
            doc.add_paragraph(request.conclusao)
        
        doc.add_paragraph()
        if request.data_extenso:
            p = doc.add_paragraph(f"Porto Velho, {request.data_extenso}.")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph(); doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(request.afte)
        run.bold, run.font.size = True, Pt(11)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Auditor(a) Fiscal de Tributos Estaduais").font.size = Pt(11)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Matrícula: {request.matricula}").font.size = Pt(11)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(request.orgao)
        run.font.size, run.font.color.rgb = Pt(9), RGBColor(128, 128, 128)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(request.endereco_orgao)
        run.font.size, run.font.color.rgb = Pt(9), RGBColor(128, 128, 128)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"relatorio_{re.sub(r'[^0-9]', '', request.cnpj)}.docx"
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}"'})
    except Exception as e:
        logger.error("[timbrado_report] Erro: %s\n%s", e, traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erro ao gerar relatório: {str(e)}")


@router.post("/reports/det-notification")
async def generate_det_notification(request: DETNotificationRequest):
    """Gera notificação DET em formato HTML."""
    try:
        data_hora = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        html_content = f"""<p style="text-align: center;"><span style="font-size:11px;">
<strong>NOTIFICAÇÃO</strong><br/>
<strong>ENCERRAMENTO DE MONITORAMENTO FISCAL</strong></span></p>
<p style="margin-left: 40px; text-align: justify;"><span style="font-size:11px;">
Razão Social: <strong>{request.razao_social}</strong><br/>CNPJ: {request.cnpj}<br/>IE: {request.ie}<br/>Endereço: {request.endereco}<br/>DSF: {request.dsf}</span></p>
<p style="margin-left: 320px; text-align: justify;"><span style="font-size:11px;"><strong>Assunto</strong>: {request.assunto}</span></p>
<p style="margin-left: 40px; text-align: justify;"><span style="font-size:11px;">{request.corpo}</span></p>
<p style="margin-left: 40px; text-align: justify;"><span style="font-size:11px;">{data_hora}</span></p>
<p style="margin-left: 40px; text-align: center;"><span style="font-size:11px;"><strong>{request.afte}</strong><br/>Auditor(a) Fiscal de Tributos Estaduais<br/>Matrícula: {request.matricula}</span></p>"""
        buffer = BytesIO(html_content.encode('utf-8'))
        filename = f"notificacao_det_{re.sub(r'[^0-9]', '', request.cnpj)}.html"
        return StreamingResponse(buffer, media_type="text/html; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{filename}"'})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/reports/det-notification-txt")
async def generate_det_notification_txt(request: DETNotificationRequest):
    """Gera notificação DET em formato TXT."""
    try:
        data_hora = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        txt_content = f"""NOTIFICAÇÃO\nENCERRAMENTO DE MONITORAMENTO FISCAL\n\nRazão Social: {request.razao_social}\nCNPJ: {request.cnpj}\nIE: {request.ie}\nEndereço: {request.endereco}\nDSF: {request.dsf}\n\nAssunto: {request.assunto}\n\n{request.corpo}\n\n{data_hora}\n\n{request.afte}\nAuditor(a) Fiscal de Tributos Estaduais\nMatrícula: {request.matricula}\n"""
        buffer = BytesIO(txt_content.encode('utf-8'))
        filename = f"notificacao_det_{re.sub(r'[^0-9]', '', request.cnpj)}.txt"
        return StreamingResponse(buffer, media_type="text/plain; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{filename}"'})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/auditoria/fisconforme")
async def gerar_fisconforme(request: FisconformeRequest):
    """Gera notificação Fisconforme para um CNPJ."""
    cnpj_limpo = re.sub(r"[^0-9]", "", request.cnpj)
    if not cnpj_limpo or not validar_cnpj(cnpj_limpo):
        raise HTTPException(status_code=400, detail="CNPJ inválido")

    try:
        from dotenv import load_dotenv
        import oracledb
        load_dotenv(dotenv_path=str(_PROJETO_DIR / ".env"), override=True)
        saved_user = os.getenv("SAVED_ORACLE_USER", "").strip().strip("'").strip('"')
        if not saved_user: raise HTTPException(status_code=401, detail="Credenciais Oracle não configuradas")
        saved_password = keyring.get_password("sefin_audit_tool", saved_user)
        if not saved_password: raise HTTPException(status_code=401, detail="Senha Oracle não encontrada")

        dsn = oracledb.makedsn("exa01-scan.sefin.ro.gov.br", 1521, service_name="sefindw")
        conexao = oracledb.connect(user=saved_user, password=saved_password, dsn=dsn)
        with conexao.cursor() as cursor:
            cursor.execute("ALTER SESSION SET NLS_NUMERIC_CHARACTERS = '.,'")
            sql_file = _PROJETO_DIR / "consultas_fonte" / "dados_cadastrais.sql"
            if not sql_file.exists(): raise HTTPException(status_code=404, detail=f"Arquivo não encontrado: {sql_file}")
            sql = ler_sql(sql_file)
            cursor.execute(sql, {"CNPJ": cnpj_limpo})
            columns = [desc[0] for desc in cursor.description]
            row = cursor.fetchone()
            dados_cadastrais = dict(zip(columns, row)) if row else {}
        conexao.close()

        from core.config_loader import get_config_var
        obter_diretorios_cnpj = get_config_var('obter_diretorios_cnpj')
        _, dir_analises, dir_relatorios = obter_diretorios_cnpj(cnpj_limpo)
        
        from gerar_relatorio import gerar_relatorio_fisconforme_html
        result = gerar_relatorio_fisconforme_html(cnpj=cnpj_limpo, dir_relatorios=dir_relatorios, dir_modelos=_PROJETO_DIR / "modelos", dados_cadastrais=dados_cadastrais, nome_auditor=request.nome_auditor, matricula_auditor=request.matricula_auditor, email_auditor=request.email_auditor, orgao=request.orgao)
        
        if not result: raise HTTPException(status_code=500, detail="Erro ao gerar notificação Fisconforme")
        return {"success": True, "file": result}
    except Exception as e:
        logger.error("[fisconforme] Erro: %s\n%s", e, traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))
